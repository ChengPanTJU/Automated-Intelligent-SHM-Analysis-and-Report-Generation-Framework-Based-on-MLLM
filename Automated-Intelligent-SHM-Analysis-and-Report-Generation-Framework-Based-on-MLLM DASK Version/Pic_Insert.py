"""
将图片和数据结果插入至word模板
"""
import os
from datetime import datetime
from operator import index

import numpy as np
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import FileConfig, OutputConfig
from docx import Document
from docx.shared import Inches

file_config = FileConfig()
OutConfig = OutputConfig()
out_config = OutConfig.tasks

def fig_insert(placeholder,doc,file_path,paragraph_style='su_图表公式'):
    '''
    placeholder:占位符；
    doc：文件；
    file_path:图片路径
    '''
    for para in doc.paragraphs:
        if placeholder in para.text:
            if os.path.exists(file_path):  # 检查图片文件是否存在
                para.text = ''  # 清空标识符文本
                # 在段落中插入图片
                run = para.add_run()
                run.add_picture(file_path)
                # 设置该段落的对齐方式为居中
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                para.style = paragraph_style
            else:
                print(f"File not found: {file_path}")  # 输出文件不存在的消息
            break  # 跳出循环，处理下一个任务

def fig_insert_new(placeholder,doc,file_path,paragraph_style='su_图表公式'):
    """
    在 Word 文档中找到指定占位符，并在其位置插入图片（不创建新段落）。
    :param placeholder: 要替换的占位符（字符串）
    :param doc: Word 文档对象 (docx.Document)
    :param file_path: 要插入的图片路径（字符串）
    :param paragraph_style: 可选，指定段落的样式，默认 "su_图表公式"
    """

    for para in doc.paragraphs:
        if placeholder in para.text:
            for run in para.runs:
                if placeholder in run.text:
                    # 拆分占位符前后的文本
                    before_text = run.text.split(placeholder)[0]
                    after_text = run.text.split(placeholder)[1]
                    # 清空原 run 文本并保留前部分
                    run.text = before_text
                    # 在当前位置插入图片
                    if not os.path.exists(file_path):
                        print(f"图片 {file_path} 不存在，无法插入")
                        para.add_run(after_text)
                        # 设置段落样式
                        para.style = paragraph_style
                        return
                    else:
                        new_run = para.add_run()
                        new_run.add_picture(file_path)  # 图片宽度可调整
                        # 恢复占位符后面的文本
                        para.add_run(after_text)
                        # 设置段落样式
                        para.style = paragraph_style
                        return  # 只替换第一个找到的占位符

def fig_insert_2(placeholder,doc,file_path,paragraph_style='su_图表公式'):
    """
    在 Word 文档中找到指定占位符，并在其位置插入图片（不创建新段落）。
    :param placeholder: 要替换的占位符（字符串）
    :param doc: Word 文档对象 (docx.Document)
    :param file_path: 要插入的图片路径（字符串）
    :param paragraph_style: 可选，指定段落的样式，默认 "su_图表公式"
    """
    for para in doc.paragraphs:
        if placeholder in para.text:
            # 拆分占位符前后的文本
            before_text = para.text.split(placeholder)[0]
            after_text = para.text.split(placeholder)[1]
            para.text=''
            # 清空原 run 文本并保留前部分
            para.add_run(before_text)
            #para.text = before_text
            # 在当前位置插入图片
            new_run = para.add_run()
            new_run.add_picture(file_path)  # 图片宽度可调整
            # 恢复占位符后面的文本
            para.add_run(after_text)
            # 设置段落样式
            para.style = paragraph_style
            return  # 只替换第一个找到的占位符

def table_insert(placeholder, doc, table_df, paragraph_style='su_图表公式'):
    '''
    placeholder:占位符；
    doc：文件；
    table_df:df格式的表格
    paragraph_style:指定的段落样式名称
    '''
    for para in doc.paragraphs:
        # 查找占位符 "{{table_tmp_mean}}"，如果找到，就在该位置插入表格
        if placeholder in para.text:
            # 在占位符处插入表格
            parent = para._element.getparent()
            # 创建一个新的表格（放置在段落前面）
            table = doc.add_table(rows=len(table_df) + 1, cols=len(table_df.columns))
            table.style = 'Table Grid'#添加边框
            # 设置表格的对齐方式和指定的段落样式
            for row in table.rows:
                for cell in row.cells:
                    # 设置单元格的段落对齐方式
                    cell.paragraphs[0].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT  # 可以根据需要修改对齐方式
                    # 应用指定的段落样式
                    cell.paragraphs[0].style = paragraph_style
                    # 为每个单元格添加边框
            # 填充表格的表头
            for col_idx, column in enumerate(table_df.columns):
                cell = table.cell(0, col_idx)
                cell.text = str(column)
                # 应用指定的段落样式
                cell.paragraphs[0].style = paragraph_style
                # 为每个表头单元格添加边框
            # 填充表格的内容
            for row_idx, row_data in enumerate(table_df.values):
                for col_idx, cell_data in enumerate(row_data):
                    cell = table.cell(row_idx + 1, col_idx)

                    # 数值统一保留两位小数
                    if pd.isna(cell_data):
                        text = ''
                    elif isinstance(cell_data, (int, float, np.number)):
                        text = f"{cell_data:.2f}"
                    else:
                        text = str(cell_data)

                    cell.text = text
                    cell.paragraphs[0].style = paragraph_style
                    # 为每个表格内容单元格添加边框
            # 插入表格到占位符段落的位置
            para._element.addnext(table._element)
            # 删除包含占位符的段落
            parent.remove(para._element)
            break  # 找到第一个匹配的占位符后就停止

def pic_insert():
    # 打开文档
    doc = Document('template_with_identifiers.docx')
    """ ========== 插入图片及任务汇总表 ============="""
    for item in file_config.filename_patterns:
        for task in out_config[item]:
            if task == 'word_identifier':
                continue
            if task=='preprocess':
                file_path_list = out_config[item][task]['raw_data_figure_path']
                file_path_list.append(out_config[item][task]['figure_path'])
                for i in range(len(file_path_list)):
                    file_path=file_path_list[i]
                    file_name_with_ext = os.path.basename(file_path)
                    file_name = os.path.splitext(file_name_with_ext)[0]
                    # 动态标识符
                    placeholder = '{'+f'{{{file_name}}}'+'}'  # 动态生成标识符，例如 '{figure_acc_Prep}'
                    fig_insert_new(placeholder, doc, file_path)
            else:
                if 'figure_path' in out_config[item][task].keys():
                    file_path_list = out_config[item][task]['figure_path']
                    if type(file_path_list)==list:
                        for i in range(len(file_path_list)):
                            file_path=file_path_list[i]
                            file_name_with_ext = os.path.basename(file_path)
                            file_name = os.path.splitext(file_name_with_ext)[0]
                            # 动态标识符
                            placeholder = '{'+f'{{{file_name}}}'+'}'  # 动态生成标识符，例如 '{figure_acc_Prep}'
                            fig_insert_new(placeholder, doc, file_path)
                    if type(file_path_list) == str:################针对traffic只有一张图的情况
                        file_path = file_path_list
                        file_name_with_ext = os.path.basename(file_path)
                        file_name = os.path.splitext(file_name_with_ext)[0]
                        # 动态标识符
                        placeholder = '{' + f'{{{file_name}}}' + '}'  # 动态生成标识符，例如 '{figure_acc_Prep}'
                        fig_insert_new(placeholder, doc, file_path)
                if 'table_identifier' in out_config[item][task].keys():
                    placeholder = out_config[item][task]['table_identifier']
                    placeholder ='{'+f'{{{placeholder}}}'+'}'
                    fill_data = pd.read_csv(out_config[item][task]['sum_table_path'])
                    table_insert(placeholder, doc, fill_data)

    ################################生成传感器情况表并插入########################################
    # 生成传感器运行状态评估表
    start_time = file_config.start_time
    end_time = file_config.end_time
    # 定义时间格式
    time_format = "%Y-%m-%d %H"
    # 定义两个时间点
    time1 = datetime.strptime(start_time, time_format)
    time2 = datetime.strptime(end_time, time_format)
    # 计算时间差
    time_difference = time2 - time1
    # 获取小时数
    hours = time_difference.total_seconds() / 3600
    columns = ['传感器类型', '编号', '正常工作时长/h','总工作时长/h', '效率比', '评估等级']
    # 创建一个空的 DataFrame，指定列名
    state_evaluate_df = pd.DataFrame(columns=columns)
    infor_dic={
        'vibration': ['主梁加速度传感器','ACC'],
        'temperature': ['温度传感器','TMP'],
        'strain': ['应变传感器','STR'],
        'traffic': ['交通流','?'],
        'displacement': ['位移传感器','DIS'],
        'cable_vib':['拉索加速度传感器','VIC'],
        'wind_speed': ['风速传感器', 'WSD'],
        'wind_direction': ['风向传感器', 'WDR'],
        'GPS': ['GPS传感器', 'GPS'],
        'inclination': ['倾角传感器', 'IAN'],
        'settlement': ['倾角传感器', 'SET'],
        'cable_force': ['索力传感器', 'CBF'],
        'humidity': ['湿度传感器', 'HUM'],
    }
    for item in file_config.filename_patterns:
        if item not in['correlation','traffic','assessment1','assessment2']:
            new_rows = []
            df=pd.read_csv(out_config[item]['preprocess']['save_path'],header=None)
            df=df.iloc[0,:]
            for i in range(len(df)):
                channel_num=out_config[item]['preprocess']['channels'][i]
                eff=df.iloc[i] * 100
                if eff >=90:
                    eff_index='I'
                elif eff>=60 and eff<90:
                    eff_index='II'
                elif eff>=30 and eff<60:
                    eff_index='III'
                else:
                    eff_index='IV'
                new_row={'传感器类型':infor_dic[item][0], '编号':f"{infor_dic[item][1]}{str(channel_num).zfill(2)}",
                         '正常工作时长/h':int(hours*df.iloc[i]),'总工作时长/h':int(hours), '效率比':str(round(df.iloc[i]*100,2))+"%", '评估等级':eff_index}
                new_rows.append(new_row)
            df_new = pd.DataFrame(new_rows,index=None)
            state_evaluate_df = pd.concat([state_evaluate_df, df_new], ignore_index=True)
    # 合并同一“传感器类型”值的单元格（逻辑上）
    state_evaluate_df['传感器类型'] = state_evaluate_df['传感器类型'].where(state_evaluate_df['传感器类型'] != state_evaluate_df['传感器类型'].shift())
    state_evaluate_df = state_evaluate_df.infer_objects(copy=False)

    # 只对“真正的字符串列”填充
    str_cols = state_evaluate_df.select_dtypes(include=["object", "string"]).columns
    state_evaluate_df[str_cols] = state_evaluate_df[str_cols].fillna(' ')
    table_insert('{{table_sensors_evaluate}}', doc, state_evaluate_df)
################################生成评分表并插入##################################
    new_rows=[]
    sensor_num_list=[]
    sensor_score_list = []
    sensor_type=0
    for item in file_config.filename_patterns:
        if item not in['correlation','traffic','assessment1','assessment2']:
            df=pd.read_csv(out_config[item]['preprocess']['save_path'],header=None)
            df=df.iloc[0,:]
            sensor_num_list.append(len(df))
            sensor_score_list.append(sum(df)/len(df)*100)
            sensor_type=sensor_type+1
    used_sensor_type=sensor_type
    for item in file_config.filename_patterns:
        if item not in['correlation','traffic','assessment1','assessment2']:
            new_row={'传感器类型':infor_dic[item][0],
                     '数量':int(sensor_num_list[sensor_type-used_sensor_type]),
                     '权重':sensor_num_list[sensor_type-used_sensor_type]/sum(sensor_num_list),
                     '得分':round(sensor_score_list[sensor_type-used_sensor_type],2)}
            new_rows.append(new_row)
        used_sensor_type=used_sensor_type-1
    new_row = {'传感器类型': "",
              '数量': "",
              '权重': "合计",
              '得分': round(sum([new_rows[i]['权重']*new_rows[i]['得分'] for i in range(len(new_rows))]),2)}
    new_rows.append(new_row)
    score_df = pd.DataFrame(new_rows,index=None)
    score_df['权重'] = score_df['权重'].apply(lambda x: round(x, 3) if isinstance(x, (int, float)) else x)
    score_df = score_df.fillna(' ')
    table_insert('{{table_sensors_score}}', doc, score_df)
    """ ========== 插入文字 ============="""
    # 保存修改后的文档
    doc.save(os.path.join('word_result','pic_result.docx'))
    print('文档已生成!')

if __name__ == '__main__':
    pic_insert()
