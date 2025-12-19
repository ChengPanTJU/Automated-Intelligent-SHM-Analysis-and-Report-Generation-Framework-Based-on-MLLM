"""
将图片和数据结果插入至word模板
"""
import os
from datetime import datetime
from operator import index

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from sympy.codegen import Print

from config import FileConfig, OutputConfig
from docx import Document
from docx.shared import Inches

file_config = FileConfig()
OutConfig = OutputConfig()
out_config = OutConfig.tasks

import os
import re
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from typing import Optional

def insert_paragraph_after(paragraph: Paragraph, text: str = "",  style: Optional[str] = None) -> Paragraph:
    """
    在指定段落后插入一个新段落，并返回新段落对象。
    """
    p = paragraph._p
    new_p = OxmlElement("w:p")
    p.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    if style:
        new_para.style = style
    return new_para


def word_insert_text(
    placeholder: str,
    doc: Document,
    file_path: str,
    paragraph_style: str = "正文A",
    delete_symbols: str = "*#-",
    collapse_empty_lines: bool = True,
    align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT,
):
    """
    将文本文件内容插入到 Word 文档中，用“真回车”（新段落）表示换行。

    功能：
      - 每个换行 -> 新段落（等同按 Enter）
      - 全文删除指定符号及其后续空格（如 "*#-" + 若干空格）
      - 连续空行折叠为 1 个空行（可关闭）

    参数：
      placeholder         用于定位的占位符文本（包含在某段落的文本中）
      doc                 python-docx 的 Document 对象
      file_path           待插入的 UTF-8 文本文件路径
      paragraph_style     段落样式名（需在模板中已存在）
      delete_symbols      要删除的符号集合字符串，默认 "*#-"
      collapse_empty_lines是否将连续空行折叠为一个空行，默认 True
      align               段落对齐方式，默认左对齐
    """
    # 1) 定位占位符所在段落
    target_para: Optional[Paragraph] = None
    for para in doc.paragraphs:
        if placeholder in para.text:
            target_para = para
            break

    if target_para is None:
        print(f"[word_insert_text] 占位符未找到：{placeholder}")
        return

    # 2) 读取文件
    if not os.path.exists(file_path):
        print(f"[word_insert_text] 文件不存在：{file_path}")
        return

    with open(file_path, "r", encoding="utf-8") as f:
        text_content = f.read()

    # 3) 统一换行符为 \n
    text_content = text_content.replace("  ", "")
    text_content = text_content.replace("\n\n", "\n").replace("\r", "\n")

    # 4) 全局删除指定符号及其后续空白
    #    例如：'--  文本' -> '文本'；'*标题' -> '标题'；'a - b' -> 'a b'
    if delete_symbols:
        sym = re.escape(delete_symbols)  # 例如 "*#-"
        text_content = re.sub(rf"[{sym}]+\s*", "", text_content)

    # 5) 折叠连续空行为 1 个空行（可选）
    if collapse_empty_lines:
        # 将 >=2 个连续（可含空白）的换行折叠为恰好两个换行（即一个空行）
        text_content = re.sub(r"(?:\n[ \t]*){2,}", "\n\n", text_content)

    # 6) 按行拆分，保留空行以生成空白段落
    lines = text_content.split("\n")

    # 7) 用第一行覆盖占位符段落文本
    target_para.text = lines[0] if lines else ""
    target_para.alignment = align
    if paragraph_style:
        target_para.style = paragraph_style

    # 8) 其余行逐行在其后插入“新段落”
    anchor = target_para
    for line in lines[1:]:
        anchor = insert_paragraph_after(anchor, line, style=paragraph_style)
        anchor.alignment = align



def word_insert():
    # 打开文档
    doc = Document(os.path.join('word_result','pic_result.docx'))
    """ ========== 插入图片及任务汇总表 ============="""
    for item in file_config.filename_patterns:
        if item not in ['correlation','assessment2']:
            s=out_config[item]
            placeholder =out_config[item]['word_identifier']
            placeholder = '{'+f'{{{placeholder}}}'+'}'
            path=os.path.join('PostProcessing', 'LLM_result',item)
            sum_files = [os.path.join(path, f) for f in os.listdir(path) if f.lower().endswith('.txt') and 'sum' in f.lower()]
            word_insert_text(placeholder, doc, sum_files[0])
        elif item =='assessment2':
            placeholder =out_config[item]['word_identifier']
            placeholder = '{'+f'{{{placeholder}}}'+'}'
            path=os.path.join('PostProcessing', 'LLM_result',item)
            sum_files = [os.path.join(path, f) for f in os.listdir(path) if f.lower().endswith('.txt') and 'reg_sum' in f.lower()]
            word_insert_text(placeholder, doc, sum_files[0])
        else:
            placeholder =out_config[item]['word_identifier']
            placeholder = '{'+f'{{{placeholder}}}'+'}'
            path=os.path.join('PostProcessing', 'LLM_result',item)
            sum_files = [os.path.join(path, f) for f in os.listdir(path) if f.lower().endswith('.txt') and 'cor_sum' in f.lower()]
            word_insert_text(placeholder, doc, sum_files[0])
    placeholder='all_summary'
    placeholder = '{' + f'{{{placeholder}}}' + '}'
    file_name=os.path.join('PostProcessing', 'LLM_result','all_sum.txt')
    word_insert_text(placeholder, doc, file_name)
    """ ========== 插入文字 ============="""
    # 保存修改后的文档
    doc.save(os.path.join('word_result',file_config.word_file_name))
    print('文档已生成!')

if __name__ == '__main__':

    word_insert()
