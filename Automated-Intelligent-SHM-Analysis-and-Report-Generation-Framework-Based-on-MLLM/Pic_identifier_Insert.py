"""
将图片和数据结果插入至word模板
"""

import os
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from config import FileConfig, OutputConfig

file_config=FileConfig()
OutConfig = OutputConfig()
out_config = OutConfig.tasks

def identifier_insert2word_new(placeholder, doc, filling_list):
    # 遍历文档中的段落
    for para in doc.paragraphs:
        if placeholder in para.text:
            # 将段落中的占位符替换为填充内容
            # 插入填充内容
            filling_text = ''.join(filling_list)  # 将列表合并为一个字符串
            para.text =filling_text  # 替换文本
            break  # 找到占位符并替换后，跳出循环

def identifier_insert2word4corr(placeholder, doc, filling_list):
    # 遍历文档中的段落
    for para in doc.paragraphs:
        if placeholder in para.text:
            for i in range(len(filling_list)):
                if '\n' in filling_list[i] or '子图' in filling_list[i]:
                    filling_list[i]=filling_list[i][2:-2]
            filling_text = ''.join(filling_list)  # 将列表合并为一个字符串
            para.text =filling_text  # 替换文本
            break  # 找到占位符并替换后，跳出循环###############

def identifier_insert():
    # 打开文档
    doc = Document('template.docx')
    # 加速度分析结果
    for item in file_config.filename_patterns:
        if item=='correlation' or item=='assessment1'or item=='assessment2':
            filling_list=[]
            for task in out_config[item]:
                if task=='word_identifier':
                    continue
                if 'figure_identifier' not in task:
                    filling_list=filling_list+out_config[item][task]['figure_path']
                    filling_list.append('\n')
                    filling_list.append(out_config[item][task]['figure_name'])
                    filling_list.append('\n')
            filling_list.pop()
            identifier = out_config[item]['figure_identifier']['figure_identifier']
            # 动态标识符
            placeholder = f'{{{identifier}}}'  # 动态生成标识符，例如 '{figure_acc_Prep}'
            filling_list = ['{' + f'{{{os.path.splitext(os.path.basename(i))[0]}}}' + '}'  for i in
                            filling_list]
            identifier_insert2word4corr(placeholder, doc, filling_list)
        elif item=='traffic':
            for task in out_config[item]:
                if task=='word_identifier':
                    continue
                if type(out_config[item][task]['figure_path']) is list:
                    identifier = out_config[item][task]['figure_identifier']
                    placeholder = f'{{{identifier}}}'  # 动态生成标识符，例如 '{figure_acc_Prep}'
                    filling_list = ['{' + f'{{{os.path.splitext(os.path.basename(i))[0]}}}' + '}' for i in
                                    out_config[item][task]['figure_path']]
                    identifier_insert2word_new(placeholder, doc, filling_list)
                else:
                    identifier = out_config[item][task]['figure_identifier']
                    placeholder = '{'+f'{{{os.path.splitext(os.path.basename(identifier))[0]}}}'+'}'
                    figname=os.path.splitext(os.path.basename(out_config[item][task]['figure_path']))[0]
                    filling_list =[ '{'+f'{{{figname}}}'+'}']
                    identifier_insert2word_new(placeholder, doc, filling_list)
        else:
            for task in out_config[item]:
                if task=='preprocess':
                    identifier = out_config[item][task]['figure_identifier']
                    placeholder = f'{{{identifier}}}'  # 动态生成标识符，例如 '{figure_acc_Prep}'
                    filling_list=[ '{'+f'{{{os.path.splitext(os.path.basename(i))[0]}}}'+'}' for i in out_config[item][task]['raw_data_figure_path']]
                    identifier_insert2word_new(placeholder, doc, filling_list)
                    identifier = out_config[item][task]['figure_path']
                    placeholder = '{'+f'{{{os.path.splitext(os.path.basename(identifier))[0]}}}'+'}'
                    filling_list =[ '{'+f'{{{os.path.splitext(os.path.basename(identifier))[0]}}}'+'}']
                    identifier_insert2word_new(placeholder, doc, filling_list)
                elif task=='word_identifier':
                    continue
                else:
                    identifier = out_config[item][task]['figure_identifier']
                    # 动态标识符
                    placeholder = f'{{{identifier}}}'  # 动态生成标识符，例如 '{figure_acc_Prep}'
                    filling_list=[ '{'+f'{{{os.path.splitext(os.path.basename(i))[0]}}}'+'}' for i in out_config[item][task]['figure_path']]
                    identifier_insert2word_new(placeholder, doc, filling_list)

    doc.save('template_with_identifiers.docx')

if __name__ == '__main__':
    file_config = FileConfig()
    OutConfig = OutputConfig()
    out_config = OutConfig.tasks
    identifier_insert()
